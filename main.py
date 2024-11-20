from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_document(input_file):
    # Create a new Document
    doc = Document()

    # Set IEEE Margins (1 inch on all sides)
    section = doc.sections[0]
    section.top_margin = 1 * 72  # 1 inch in points
    section.bottom_margin = 1 * 72
    section.left_margin = 1 * 72
    section.right_margin = 1 * 72

    # Add Title Page in IEEE format
    add_title_page(doc)

    # Read the input text file and convert it into DOCX
    with open(input_file, 'r') as file:
        lines = file.readlines()

    for line in lines:
        # Handle Heading (main or sub-headings)
        # Main Section Heading (e.g., 1, 2, 3)
        first=line.split(" ")
        if first[0].count(".")==1:
            add_main_section_heading(doc, line.strip())
        # Subsection Heading (e.g., 1.1, 2.2)
        elif first[0].count(".")==2:
            add_subsection_heading(doc, line.strip())
        # Sub-subsection Heading (e.g., 1.1.1, 2.2.2)
        elif first[0].count(".")==3:
            add_subsubsection_heading(doc, line.strip())
        else:  # Regular body text
            add_paragraph(doc, line.strip())

    # Save the document to a file
    doc.save('output_ieee.docx')

def add_title_page(doc):
    # Title Page Structure (IEEE format)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Title of the Paper")
    title_run.bold = True
    title_run.font.size = Pt(24)  # Title size

    doc.add_paragraph()  # Add space after title

    # Authors and Affiliations
    authors = doc.add_paragraph()
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    authors_run = authors.add_run("Author 1, Author 2, Author 3")
    authors_run.font.size = Pt(12)  # Author size

    doc.add_paragraph()  # Add space after authors

    # Author Affiliations and Emails
    affiliations = doc.add_paragraph()
    affiliations.alignment = WD_ALIGN_PARAGRAPH.CENTER
    affiliations_run = affiliations.add_run("Affiliation 1, Email 1\nAffiliation 2, Email 2")
    affiliations_run.font.size = Pt(10)  # Smaller font for affiliations

    # Add space before abstract
    doc.add_paragraph() 

    # Abstract
    abstract = doc.add_paragraph("Abstract: This is the abstract of the paper.")
    abstract.alignment = WD_ALIGN_PARAGRAPH.LEFT
    abstract_run = abstract.add_run()
    abstract_run.font.size = Pt(12)

    doc.add_paragraph()  # Add space after abstract

def add_main_section_heading(doc, text):
    # Main Section Heading (IEEE numbered, 14pt, bold)
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    heading_run = heading.add_run(text)
    heading_run.bold = True
    heading_run.font.size = Pt(14)  # Section Heading size

def add_subsection_heading(doc, text):
    # Subsection Heading (IEEE, 12pt, bold)
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    heading_run = heading.add_run(text)
    heading_run.bold = True
    heading_run.font.size = Pt(12)  # Subsection Heading size

def add_subsubsection_heading(doc, text):
    # Sub-subsection Heading (IEEE, 12pt, italicized)
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    heading_run = heading.add_run(text)
    heading_run.italic = True
    heading_run.font.size = Pt(12)  # Sub-subsection Heading size

def add_paragraph(doc, text):
    # Add regular paragraph with specific formatting
    paragraph = doc.add_paragraph(text)
    paragraph.style.font.size = Pt(12)  # Font size 12pt (IEEE body text)
    paragraph.paragraph_format.line_spacing = 1  # Single line spacing for IEEE

if __name__ == "__main__":
    input_file = "sample.txt"  
    create_document(input_file)
